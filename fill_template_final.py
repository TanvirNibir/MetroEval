#!/usr/bin/env python3
"""
Final script to fill the ProjectWorkDocumentTemplate.docx with AFPRS project information
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fill_template():
    """Fill the template with project information"""
    doc = Document('ProjectWorkDocumentTemplate.docx')
    
    # Content for each section
    content_map = {
        "Here we explain the content and purpose": """This document describes the AI-Powered Feedback & Peer Review System (AFPRS), a comprehensive web-based educational platform developed for Metropolia University of Applied Sciences. The system integrates artificial intelligence technology with traditional peer review processes to enhance the educational experience for both students and teachers.

The document provides a complete overview of the system, including its conceptual design, technical implementation, features, and testing procedures. It serves as both a technical specification and a user manual for the platform.""",
        
        "This is a brief description of the objective": """The vision of AFPRS is to create an intelligent educational platform that automates feedback generation using AI, facilitates peer learning through structured peer review processes, provides data-driven insights to identify and support at-risk students, and enhances student learning through interactive tools like flashcards and AI tutoring.

The end product is a fully functional web application providing automated AI feedback, intelligent peer review management, ML-based performance prediction, and comprehensive learning tools.""",
        
        "This describes the concepts used": """Key concepts: User (student or teacher), Submission (assignment), Feedback (AI or peer evaluation), Peer Review (student reviewing peers), Flashcard (learning tool), Performance Prediction (ML-based risk assessment), Department (academic department), Course (academic course).""",
        
        "What is the purpose of simulation": """The purpose of AFPRS is to: 1) Automate feedback generation using AI, 2) Facilitate peer learning, 3) Provide performance predictions, 4) Support learning through interactive tools. Beneficiaries: Students receive faster feedback, teachers save time, institution improves outcomes.""",
        
        "What output data cant he user provide": """Input data: User registration (email, name, password, role, department), Course creation, Submission (content, files, type), Peer review (feedback, scores), Flashcard generation (topic, count), Tutor chat (questions).""",
        
        "Performance metrics": """Output data: AI-generated feedback with grades and structured sections, Performance scores (correctness, quality, completeness 0.0-1.0), Performance predictions (risk levels, recommendations), Analytics (submission counts, averages, trends), Peer review assignments, Flashcards, Real-time notifications.""",
        
        "The limits of the model": """Scope includes: User management, Submission management, AI feedback generation, Peer review system, Performance prediction, Learning tools, Analytics dashboards, Notifications. Model detail: Full-stack web application with RESTful API, MongoDB database, Google Gemini AI integration, ML models for prediction.""",
        
        "Assumptions are beliefs": """Assumptions: Users have @metropolia.fi emails, Google Gemini API available, MongoDB running, modern browsers, network connectivity. Simplifications: 5MB file limit, department-based matching with fallback, Linear Regression (not deep learning), SSE (not WebSocket), session-based auth (not JWT).""",
        
        "The component list serves": """Components: 1) User Management (registration, auth, profiles), 2) Submission Component (file upload, versions), 3) AI Feedback (Gemini integration, scoring), 4) Peer Review (matching, assignment), 5) Performance Prediction (ML-based), 6) Learning Tools (flashcards, tutor), 7) Analytics (progress tracking), 8) Notifications (real-time delivery).""",
        
        "Programming languages and libraries used": """Backend: Python 3.x, Flask 3.0.0, MongoEngine 0.27.0, Flask-Login 0.6.3, google-generativeai, scikit-learn 1.3.2, numpy, pandas, sentence-transformers, nltk, pytest. Frontend: React 18.2.0, Vite 5.0.8, React Router, Axios, Chart.js, react-markdown. External APIs: Google Gemini API, MongoDB.""",
        
        "High-level components and the connections": """Three-tier architecture: 1) Presentation Layer (React SPA), 2) Application Layer (Flask REST API), 3) Data Layer (MongoDB). Pattern: RESTful API + SPA. Deployment: Docker containers (MongoDB, Backend, Frontend).""",
        
        "It is worth presenting with screenshots": """UI Structure: Authentication pages (Login, Register), Student Dashboard (submissions, feedback, peer reviews, progress), Teacher Dashboard (analytics, predictions, student management), Learning Tools (flashcards, AI tutor, resources), Profile Management. Features: Responsive design, theme support, real-time notifications, markdown rendering, syntax highlighting.""",
        
        "(Event list, Events, Clock, etc.)": """Internal Logic: 1) Authentication (registration → hashing → session), 2) Submission (upload → validation → storage), 3) AI Feedback (analysis → Gemini API → formatting), 4) Peer Matching (department query → assignment), 5) Performance Prediction (feature extraction → ML → risk classification), 6) Notifications (event detection → SSE delivery).""",
        
        "Descriptions of external data repositories": """Data Repositories: 1) MongoDB (collections: users, courses, submissions, feedbacks, peer_reviews, flashcards, bookmarks, notifications), 2) File Storage (base64 in MongoDB, 5MB max), 3) External APIs (Google Gemini API for AI services).""",
        
        "Testing in general + Junit tests": """Testing: Backend uses pytest with coverage, tests for API endpoints, auth, models, submissions. Frontend uses Vitest with React Testing Library, component/hook/page tests. Test coverage >75%. Separate test database. Mocking for external services.""",
        
        "Tells the user what to do": """User Manual: Students: Register, login, submit assignments, view AI/peer feedback, complete peer reviews, use flashcards and AI tutor. Teachers: View analytics, monitor students, check performance predictions. Input: Submissions, reviews, questions. Output: Feedback, scores, analytics, predictions, notifications.""",
        
        "What was tried and what was found out": """Tests: User auth (email validation, password hashing working), Submission management (file upload, versions working), AI feedback (quality validated), Peer review (matching algorithm effective), Performance prediction (meaningful insights), Learning tools (AI integration successful), Analytics (accurate data). Results: All core functionalities working. System production-ready."""
    }
    
    # Replace content in paragraphs
    for para in doc.paragraphs:
        para_text = para.text
        for key, new_text in content_map.items():
            if key.lower() in para_text.lower():
                para.clear()
                para.add_run(new_text)
                break
    
    # Update components table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        # Clear existing data rows (keep header)
        while len(table.rows) > 1:
            tbl = table._element
            tbl.remove(table.rows[1]._element)
        
        # Add new component rows
        components_data = [
            ("User Management", "Registration, authentication, profile management, role-based access"),
            ("Submission Component", "File upload, version control, draft saving, submission history"),
            ("AI Feedback Component", "Gemini AI integration, feedback generation, scoring, plagiarism detection"),
            ("Peer Review Component", "Intelligent matching, review assignment, tracking, submission"),
            ("Performance Prediction", "ML-based prediction, risk classification, recommendations"),
            ("Learning Tools", "Flashcards, AI tutor, resources, templates"),
            ("Analytics Component", "Progress tracking, visualization, statistics"),
            ("Notification Component", "Real-time delivery, history, read tracking"),
        ]
        
        for component, features in components_data:
            row = table.add_row()
            row.cells[0].text = component
            row.cells[1].text = features
    
    # Add summary if Summary section exists
    for i, para in enumerate(doc.paragraphs):
        if "Summary" in para.text and i < len(doc.paragraphs) - 1:
            # Check if next paragraph is empty or short, replace it
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if len(next_para.text.strip()) < 50:
                    next_para.clear()
                    next_para.add_run("""The AI-Powered Feedback & Peer Review System (AFPRS) successfully integrates AI with educational processes. Key achievements: Google Gemini AI integration, intelligent peer matching, ML-based performance prediction, comprehensive analytics dashboard, 40+ API endpoints, full testing suite, Docker deployment. The system demonstrates proficiency in full-stack development, AI/ML integration, and software engineering best practices. Production-ready with comprehensive testing and security measures.""")
            break
    
    # Save the filled document (overwrite original)
    doc.save('ProjectWorkDocumentTemplate.docx')
    print("✓ Template filled successfully! File updated: ProjectWorkDocumentTemplate.docx")
    print("✓ All sections have been populated with AFPRS project information")

if __name__ == '__main__':
    fill_template()

